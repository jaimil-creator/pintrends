import os
import json
import re
import uuid
from typing import List, Dict, Tuple
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from pathlib import Path
from dotenv import load_dotenv

# Load env from root
env_path = Path(__file__).resolve().parent.parent.parent / '.env'
load_dotenv(dotenv_path=env_path)

# ==================== SYSTEM PROMPTS ====================

BLOG_SYSTEM_PROMPT = """You are a trendy, energetic 18-year-old Pinterest blog writer. Your goal is to write a high-engagement blog post that feels like a conversation on Reddit or a viral social media caption.

The topic is: {topic}

STRICT CONSTRAINT ENFORCEMENT (HIGHEST PRIORITY):
1. Identify the ANCHOR ITEM: (e.g., If topic is "Black Mini Skirt Outfits", the Anchor is "Black Mini Skirt").
2. Title Consistency: Every single Idea Title MUST mention or focus on the ANCHOR ITEM.
   - BAD: "Oversized Cream Knit" (Focuses on the sweater, not the skirt).
   - GOOD: "Leather Mini with Cream Knit" (Focuses on the skirt pairing).
3. Color & Item Logic: The Anchor Item must ALWAYS match the requested color/type.
   - If topic is "Black Skirts", never describe a "Red Skirt".
   - You can vary the *fabric* (Leather, Denim, Velvet, Pleated) but the item must remain a [Color] [Item].

CRITICAL VOICE & STYLE RULES:

Voice: Direct, simple, and energetic. Use confident statements.

Tone: Warm, cozy, casual, modern, and no robotic tone.

Language: Use simple words. No complicated vocabulary. Write like a real person, not a textbook.

Formatting: ABSOLUTELY NO DASHES (— or -) inside sentences. Use periods or commas only. No fluff.

Paragraph Structure: Descriptions should usually be one solid paragraph, but randomly split some descriptions into 2 shorter paragraphs to make the text look good and easy to read. Do not do this for every item; keep it random.

No Lists Inside Descriptions: Do not use bullet points or numbered lists within the paragraph descriptions.

Topic Relevance: The content must be 100% focused on the requested topic. Do not drift into general advice.

SEO & Keywords: You must mention the exact topic keyword naturally in the Intro and occasionally in the descriptions. Do not force it, but make sure the reader knows exactly what this post is about.

Zero Plagiarism: 100% original writing.

SUBHEADING CONTENT ALIGNMENT RULES:

Subsection Headings should directly match the description. If the heading is “Leather Mini with Boots,” the description must focus on that specific skirt look.

Avoid Repetitiveness in Starting Phrases: Don’t always begin descriptions with the same structure. Avoid using repetitive starting phrases like "Go with" or "Pick a" in every description. Mix up the sentence structure to keep it fresh and engaging.

NEGATIVE WORD LIST (STRICTLY FORBIDDEN):

Do NOT use any of the following words or phrases:
Transform, unveil the secrets, Enhance, robust, Meticulous, explore, Navigating, discover, Complexities, Realm, welcome to world of, Bespoke, captivate, Tailored, breathtaking, Towards, traits, Underpins, in the heart of, Ever-changing, when it comes to, Ever-evolving, in the realm of, The world of, Not only, unlock the secrets, Seeking more than just, our suite, Designed to enhance, it is advisable, You know those mornings, just hit different, Let’s be real for a second, Whether you are looking, There is something, The holiday season is finally here.

BLOG STRUCTURE:

1. Intro (Medium length)
Write a direct, descriptive intro that immediately addresses the specific topic.

Hook: Start with a hook 1-2 sentence and then start to relate directly to the {topic}.

Word count: 100-120 words

Keyword Usage: You must include the main keyword phrase naturally.

Content: Connect the mood directly to the fashion/topic.

Ending: Must include a short Call to Action (CTA) at the end of the intro.

2. Numbered Ideas
Generate a list of ideas related to the topic. If a number is not specified in the input, pick 5-7 strong ideas.

Idea Title: Short, fun, and simple. MUST include the Main Topic Item or a variation of it (e.g. mention the fabric/cut of the skirt).

Idea Description (4-5 lines):

Content: Talk ONLY about the specific subheading. Describe the specific outfit but keep the focus on the MAIN TOPIC ITEM.

Grammar: Simple sentences. Correct grammar. NO DASHES.

Formatting: Randomly choose to keep it as 1 paragraph or split it into 2 paragraphs for visual variety.

3. Conclusion (Short & Punchy)

Wrap up the specific topic.

The Goal: Remind them why these specific ideas are worth trying.

The Ending: Finish with a Single, Strong Call to Action (CTA) (e.g., "Save these ideas to Pinterest now" or "Click the link to shop").

FINAL INSTRUCTION:

Only output the blog post content (Intro, Ideas, Conclusion). Do not generate a main Title. Do not add conversational filler.

REQUIRED OUTPUT FORMAT:
---INTRO---
(Intro content here)

---IDEAS---
1. (Idea Title)
(Idea Description)

...

---CONCLUSION---
(Conclusion content here)
"""

SYSTEM_MESSAGE_THUMBNAILS = '''Create an image generation prompt for a blog thumbnail.

INPUT FORMAT:
Main Blog Title: [text]
Section Title: [text]
Description: [text]

STEP 1 - EXTRACT ATTRIBUTES FROM MAIN BLOG TITLE:
Look for these optional attributes in the Main Blog Title and include them if found:

MODEL ATTRIBUTES:
- Gender: men, women, male, female
- Height: tall, short, petite, average height
- Body shape/size: curvy, plus-size, slim, athletic, hourglass, pear-shaped, apple-shaped, rectangle
- Hair: blonde, brunette, redhead, black hair, curly hair, straight hair, short hair, long hair, wavy hair
- Skin tone: dark skin, brown skin, fair skin, light skin, tan, olive skin, pale
- Age: young, mature, teenage, 20s, 30s, 40s, 50s, over 50, elderly

CONTEXT ATTRIBUTES:
- Season & Weather: spring, summer, fall/autumn, winter, rainy, sunny, snowy, cloudy, warm, cold, humid
- Occasion: casual, formal, office, wedding, party, date night, beach, gym, travel, vacation, brunch, festival
- Location: urban, city, street, beach, countryside, mountains, cafe, office, home, garden, rooftop, mall
- Fashion Aesthetics & Trends: minimalist, bohemian, streetwear, vintage, Y2K, cottagecore, dark academia, old money, coastal grandmother, quiet luxury, preppy, grunge, romantic, edgy
- Budget & Shopping Style: affordable, luxury, thrift, designer, budget-friendly, high-end, sustainable
- Material & Fabric: cotton, linen, silk, velvet, leather, denim, wool, cashmere, satin, chiffon, knit, crochet
- Color: neutral, pastel, bold, monochrome, earth tones, jewel tones, black, white, beige, navy, red, pink, green

STEP 2 - CREATE A UNIQUE BACKGROUND:
CRITICAL: Each image MUST have a completely unique, natural-looking background that:
1. Looks like a real candid street-style or lifestyle photograph - NOT staged or artificial
2. Uses natural outdoor environments: city streets, parks, beaches, countryside, cafes, real buildings
3. Complements the clothes with natural color harmony and textures
4. Feels organic and authentic - like a real fashion photographer captured it on location
5. Includes natural light and real-world environmental elements

AVOID: Studio setups, artificial lighting, overly polished or staged environments, green screens, or anything that looks fake.

PREFER: Real streets, natural parks, authentic cafes and restaurants, genuine urban environments, natural landscapes, real architectural spaces with character.

STEP 3 - BUILD THE PROMPT:
Include all extracted attributes naturally with your unique background.

OUTPUT FORMAT:
"Create a realistic editorial photograph with NO text, showing [model with attributes] wearing [outfit], in [YOUR UNIQUE CREATIVE BACKGROUND with specific details], [occasion/season context], [fashion aesthetic] style. [Specific lighting that matches the setting]. Pinterest aesthetic."

EXAMPLES:

Input: Main Blog Title: Office outfit ideas for women
Output: "Create a realistic editorial photograph with NO text, showing a female model in a contemporary glass-walled office lobby with marble floors and lush indoor plants, morning light streaming through floor-to-ceiling windows, professional workwear, minimalist style. Soft diffused natural lighting. Pinterest aesthetic."

Input: Main Blog Title: Winter outfits for curvy women with dark skin
Output: "Create a realistic editorial photograph with NO text, showing a curvy dark-skinned female model on a charming European cobblestone street dusted with fresh snow, vintage street lamps glowing warmly, cozy winter fashion, warm and stylish aesthetic. Golden hour winter lighting with soft snowfall. Pinterest aesthetic."

Input: Main Blog Title: Old money aesthetic fall outfits in neutral colors
Output: "Create a realistic editorial photograph with NO text, showing a model in a grand library with mahogany bookshelves and leather armchairs, autumn leaves visible through arched windows, old money aesthetic, neutral earth tones, quiet luxury style. Rich warm afternoon light. Pinterest aesthetic."

NOW: Read the input below and output ONLY the prompt. No questions, no explanations.
'''

SYSTEM_MESSAGE_IMAGES = '''Create an image generation prompt for a specific blog item.

INPUT FORMAT:
Main Blog Title: [text]
Section Title: [text]  
Description: [text]

STEP 1 - EXTRACT ATTRIBUTES FROM MAIN BLOG TITLE:
Look for these optional attributes in the Main Blog Title and include them if found:

MODEL ATTRIBUTES:
- Gender: men, women, male, female
- Height: tall, short, petite, average height
- Body shape/size: curvy, plus-size, slim, athletic, hourglass, pear-shaped, apple-shaped, rectangle
- Hair: blonde, brunette, redhead, black hair, curly hair, straight hair, short hair, long hair, wavy hair
- Skin tone: dark skin, brown skin, fair skin, light skin, tan, olive skin, pale
- Age: young, mature, teenage, 20s, 30s, 40s, 50s, over 50, elderly

CONTEXT ATTRIBUTES:
- Season & Weather: spring, summer, fall/autumn, winter, rainy, sunny, snowy, cloudy, warm, cold, humid
- Occasion: casual, formal, office, wedding, party, date night, beach, gym, travel, vacation, brunch, festival
- Location: urban, city, street, beach, countryside, mountains, cafe, office, home, garden, rooftop, mall
- Fashion Aesthetics & Trends: minimalist, bohemian, streetwear, vintage, Y2K, cottagecore, dark academia, old money, coastal grandmother, quiet luxury, preppy, grunge, romantic, edgy
- Budget & Shopping Style: affordable, luxury, thrift, designer, budget-friendly, high-end, sustainable
- Material & Fabric: cotton, linen, silk, velvet, leather, denim, wool, cashmere, satin, chiffon, knit, crochet
- Color: neutral, pastel, bold, monochrome, earth tones, jewel tones, black, white, beige, navy, red, pink, green

STEP 2 - CREATE A UNIQUE NATURAL BACKGROUND FOR THIS SPECIFIC ITEM:
CRITICAL: Each blog item image MUST have a DIFFERENT, unique background that looks completely NATURAL and REALISTIC.

Your background should:
1. Look like authentic street-style or lifestyle photography - candid and organic, NOT staged
2. Use real outdoor locations: actual city streets, parks, beaches, gardens, outdoor cafes, real neighborhoods
3. Feature natural daylight and authentic environmental elements
4. Feel like a real fashion photographer captured it on location during a natural moment
5. Include organic textures and real-world details that make it feel genuine

AVOID: Artificial setups, studio backgrounds, overly polished environments, anything that looks fake or digitally created.

PREFER: Tree-lined streets, waterfront walks, outdoor cafe patios, city sidewalks, park pathways, beach boardwalks, natural gardens, authentic urban neighborhoods, real architectural facades with character.

STEP 3 - BUILD THE PROMPT:
Use EXACT words from Section Title as the clothing/item.
Include all extracted model and context attributes naturally with your unique background.

OUTPUT FORMAT:
"Create a realistic editorial photograph with NO text showing ONLY ONE PERSON - no other people in the background or frame. A [model attributes] model wearing [EXACT Section Title], in [YOUR UNIQUE CREATIVE BACKGROUND with specific environmental details], [mood from description], [fashion aesthetic] style. [Specific lighting for this setting]. Pinterest aesthetic."

IMPORTANT RULES:
- The Section Title clothing item MUST appear exactly as written in the output
- Background MUST be unique, creative and immersive - never generic!
- Include specific details that make the background feel real and alive

EXAMPLES:

Input:
Main Blog Title: Office outfit ideas for women
Section Title: Oversized Blazer and Jeans
Output: "Create a realistic editorial photograph with NO text. A female model wearing an oversized blazer and jeans, in a sleek rooftop lounge with city skyline views and modern outdoor furniture, professional and chic, minimalist style. Late morning diffused sunlight. Pinterest aesthetic."

Input:
Main Blog Title: Winter outfits for curvy women with dark skin
Section Title: Chunky Knit Sweater
Output: "Create a realistic editorial photograph with NO text. A curvy dark-skinned female model wearing a chunky knit sweater, in a cozy Nordic-style cabin with a stone fireplace and knitted blankets, warm mugs on a wooden table, cozy and warm, winter casual style. Soft amber firelight mixed with cool window light. Pinterest aesthetic."

Input:
Main Blog Title: Affordable summer beach looks for plus-size women over 40
Section Title: Flowy Linen Maxi Dress
Output: "Create a realistic editorial photograph with NO text. A plus-size mature female model wearing a flowy linen maxi dress, on a whitewashed Mediterranean terrace with bougainvillea cascading over blue doors, relaxed and elegant, vacation style. Bright golden hour Mediterranean light. Pinterest aesthetic."

Input:
Main Blog Title: Streetwear for tall athletic men in the city
Section Title: Oversized Hoodie and Joggers
Output: "Create a realistic editorial photograph with NO text. A tall athletic male model wearing an oversized hoodie and joggers, in a gritty urban underpass with colorful street art murals and puddles reflecting neon signs, urban and cool, streetwear edgy style. Dramatic mixed artificial and natural lighting. Pinterest aesthetic."

GENDER RULE: Use "male model" for men's items, "female model" for women's items (if not specified, infer from context).

NOW: Read the input below and output ONLY the prompt. No questions, no explanations.
'''


class BlogGeneratorService:
    """
    AI-powered blog generation service using exact BLOG_GEN prompts and configuration.
    Generates complete blogs with AI-generated images using multiple AI providers.
    """
    
    def __init__(self):
        # API Keys
        self.gemini_key = os.getenv("GEMINI_API_KEY")
        self.together_key = os.getenv("TOGETHER_API_KEY")
        self.fal_key = os.getenv("FAL_KEY")
        
        # Initialize clients
        self._init_clients()
    
    def _init_clients(self):
        """Initialize AI service clients."""
        # Gemini client for blog content
        if self.gemini_key:
            try:
                from google import genai
                self.gemini_client = genai.Client(api_key=self.gemini_key)
            except Exception as e:
                print(f"Failed to initialize Gemini client: {e}")
                self.gemini_client = None
        else:
            self.gemini_client = None
        
        # Together AI client for image prompts
        if self.together_key:
            try:
                from together import Together
                self.together_client = Together(api_key=self.together_key)
            except Exception as e:
                print(f"Failed to initialize Together client: {e}")
                self.together_client = None
        else:
            self.together_client = None
        
        # Fal client for image generation
        if self.fal_key:
            try:
                import fal_client
                os.environ["FAL_KEY"] = self.fal_key  # Fal uses env var
                self.fal_available = True
            except Exception as e:
                print(f"Failed to initialize Fal client: {e}")
                self.fal_available = False
        else:
            self.fal_available = False
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((Exception,))
    )
    def generate_blog_content(self, topic: str) -> str:
        """
        Generate blog content using Google Gemini - exact BLOG_GEN implementation.
        Returns formatted blog content with ---INTRO---, ---IDEAS---, ---CONCLUSION---.
        """
        if not self.gemini_client:
            raise Exception("Gemini API key not configured. Please set GEMINI_API_KEY in .env file.")
        
        from google.genai import types
        
        formatted_system_prompt = BLOG_SYSTEM_PROMPT.replace("{topic}", topic)
        user_prompt = f"topic: {topic}"
        
        try:
            # Gemini 3 Pro supports system role in contents
            contents = [
                types.Content(
                    role="system",
                    parts=[
                        types.Part.from_text(text=formatted_system_prompt),
                    ],
                ),
                types.Content(
                    role="user",
                    parts=[
                        types.Part.from_text(text=user_prompt),
                    ],
                ),
            ]
            
            generate_content_config = types.GenerateContentConfig()
            
            print("\n📝 Generating blog content with Gemini Flash...")
            
            # Use gemini-2.0-flash-exp which has available quota
            # (gemini-3-pro-preview quota exceeded)
            full_text_parts = []
            for chunk in self.gemini_client.models.generate_content_stream(
                model='gemini-3-flash-preview',
                contents=contents,
                config=generate_content_config,
            ):
                if chunk.text:
                    full_text_parts.append(chunk.text)
            
            full_text = ''.join(full_text_parts)
            
            if not full_text or len(full_text) < 100:
                raise ValueError(f"Generated content too short ({len(full_text)} chars)")
            
            return full_text
            
        except Exception as e:
            print(f"Blog generation error: {e}")
            raise
    
    def parse_blog_content(self, content: str) -> Tuple[str, List[Dict[str, str]], str]:
        """
        Parse generated blog content - exact BLOG_GEN parsing logic.
        Returns: (intro, items_list, conclusion)
        """
        intro = ""
        items = []
        conclusion = ""
        
        # Split by main sections
        parts = content.split("---IDEAS---")
        
        # Parse Intro
        if len(parts) > 0:
            intro_section = parts[0]
            if "---INTRO---" in intro_section:
                intro = intro_section.split("---INTRO---")[1].strip()
            else:
                intro = intro_section.strip()
                
        # Parse Ideas and Conclusion
        if len(parts) > 1:
            remaining = parts[1]
            if "---CONCLUSION---" in remaining:
                ideas_part, conclusion_part = remaining.split("---CONCLUSION---")
                conclusion = conclusion_part.strip()
            else:
                ideas_part = remaining
                
            # Parse individual ideas
            lines = ideas_part.strip().split('\n')
            current_item = None
            current_description = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check for numbered item
                match = re.match(r'^(\d+)[.):\-\s]+(.+)$', line)
                
                if match:
                    # Save previous item
                    if current_item:
                        items.append({
                            'title': current_item,
                            'description': ' '.join(current_description).strip()
                        })
                    
                    # Start new item
                    current_item = match.group(2).strip()
                    current_description = []
                elif current_item:
                    current_description.append(line)
            
            # Add last item
            if current_item:
                items.append({
                    'title': current_item,
                    'description': ' '.join(current_description).strip()
                })
                
        return intro, items, conclusion
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def generate_image_prompt(self, title: str, description: str, prompt_type: str = "image", blog_topic: str = None) -> str:
        """Generate AI image prompt using Together AI with Qwen model - exact BLOG_GEN implementation."""
        if not self.together_client:
            return f"Lifestyle photography of {title}, professional editorial style, natural lighting"
        
        if prompt_type == "thumbnail":
            system_message = SYSTEM_MESSAGE_THUMBNAILS
            model_name = "Thumbnail"
        else:
            system_message = SYSTEM_MESSAGE_IMAGES
            model_name = "Blog Image"

        # Matched logic from BLOG_GEN/final_blog_generator.py
        dominant_title = blog_topic if blog_topic else title
        section_title = title # In BLOG_GEN, they use 'title' as section title
        
        user_message = f"""Main Blog Title: {dominant_title}
Section Title: {section_title}
Description: {description}"""
        
        try:
            print(f"\n🤖 Generating {model_name} prompt with Qwen...")
            print(f"   Main Blog Title: {dominant_title}")
            print(f"   Section Title: {section_title}")
            
            response = self.together_client.chat.completions.create(
                model="Qwen/Qwen3-Next-80B-A3B-Instruct", # Authentically using Qwen as per reference
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ],
                max_tokens=1000,
                temperature=0.7,
                top_p=0.9,
            )
            
            generated_prompt = response.choices[0].message.content.strip()
            generated_prompt = ' '.join(generated_prompt.split())
            
            print(f"✅ Generated prompt: {generated_prompt[:150]}...")
            
            return generated_prompt
        except Exception as e:
            print(f"Error generating {prompt_type} prompt: {e}")
            return f"Lifestyle photography of {section_title}, professional editorial style, natural lighting"
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=3, max=15)
    )
    def generate_image(self, prompt: str, aspect_ratio: str = "3:4") -> str:
        """Generate image using Fal AI with nano-banana model - exact BLOG_GEN implementation."""
        if not self.fal_available:
            raise Exception("FAL_KEY not configured. Please set FAL_KEY in .env file.")
        
        try:
            import fal_client
            
            # Convert aspect ratio format
            if aspect_ratio == "3:4":
                fal_aspect = "3:4"
            elif aspect_ratio == "16:9":
                fal_aspect = "16:9"
            else:
                fal_aspect = "3:4"
            
            print(f"🎨 Generating Image with Fal AI (Aspect Ratio: {fal_aspect})...")
            
            result = fal_client.subscribe(
                "fal-ai/nano-banana",
                arguments={
                    "prompt": prompt,
                    "num_images": 1,
                    "aspect_ratio": fal_aspect,
                    "output_format": "png"
                }
            )
            
            if result and 'images' in result and len(result['images']) > 0:
                image_url = result['images'][0]['url']
                print(f"✅ Image Generated: {image_url[:80]}...")
                return image_url
            else:
                raise Exception("No image generated")
        except Exception as e:
            print(f"Image generation error: {e}")
            raise
    
    def generate_all_images_parallel(self, prompts: Dict[str, str]) -> Dict[str, str]:
        """Generate all blog images in parallel."""
        images = {}
        
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {}
            
            for key, prompt in prompts.items():
                if prompt:
                    aspect_ratio = "16:9" if key == 'thumbnail' else "3:4"
                    futures[key] = executor.submit(self.generate_image, prompt, aspect_ratio)
            
            for key, future in futures.items():
                try:
                    images[key] = future.result(timeout=90)
                    print(f"✓ Generated image for {key}")
                except Exception as e:
                    print(f"✗ Error generating image for {key}: {e}")
                    images[key] = ""
        
        return images
    
    def download_image(self, url: str) -> BytesIO:
        """Download image from URL to BytesIO stream."""
        import requests
        
        try:
            response = requests.get(url, timeout=15)
            response.raise_for_status()
            return BytesIO(response.content)
        except Exception as e:
            print(f"Error downloading image: {e}")
            return None
    
    def create_docx(self, blog_data: Dict) -> BytesIO:
        """Create Word document from blog data with parallel image downloading."""
        from docx import Document
        from docx.shared import Inches
        
        # 1. Collect all image URLs
        urls_to_download = []
        if blog_data.get('thumbnail_url'):
            urls_to_download.append(blog_data['thumbnail_url'])
        
        for section in blog_data.get('sections', []):
            if section.get('image_url'):
                urls_to_download.append(section['image_url'])
        
        # 2. Download images in parallel
        downloaded_images = {}
        if urls_to_download:
            print(f"⬇️ Downloading {len(urls_to_download)} images for DOCX generation...")
            with ThreadPoolExecutor(max_workers=8) as executor:
                # Create a map of future -> url
                future_to_url = {executor.submit(self.download_image, url): url for url in urls_to_download}
                
                for future in future_to_url:
                    url = future_to_url[future]
                    try:
                        image_stream = future.result()
                        if image_stream:
                            downloaded_images[url] = image_stream
                    except Exception as e:
                        print(f"Failed to download image {url}: {e}")

        # 3. Create Document
        doc = Document()
        doc.add_heading(blog_data['topic'].upper(), 0)
        
        # Add thumbnail
        if blog_data.get('thumbnail_url') and blog_data['thumbnail_url'] in downloaded_images:
            try:
                # Reset stream position just in case
                img_stream = downloaded_images[blog_data['thumbnail_url']]
                img_stream.seek(0) 
                doc.add_picture(img_stream, width=Inches(6))
            except Exception as e:
                print(f"Error adding thumbnail: {e}")
        
        # Add intro
        doc.add_paragraph(blog_data['intro'])
        
        # Add sections
        for section in blog_data['sections']:
            doc.add_heading(section['title'], level=1)
            
            if section.get('image_url') and section['image_url'] in downloaded_images:
                try:
                    img_stream = downloaded_images[section['image_url']]
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(4))
                except Exception as e:
                    print(f"Error adding section image: {e}")
            
            doc.add_paragraph(section['description'])
        
        # Add conclusion
        if blog_data.get('conclusion'):
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(blog_data['conclusion'])
        
        # Save to BytesIO
        docx_stream = BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream
    
    def create_pinterest_json(self, blog_data: Dict) -> Dict:
        """Create Pinterest-compatible JSON payload."""
        return {
            "id": f"pinterest-blog-{uuid.uuid4().hex[:8]}",
            "title": blog_data['topic'],
            "thumbnail_url": blog_data.get('thumbnail_url', ''),
            "alt": f"{blog_data['topic']} thumbnail",
            "description": [blog_data['intro']],
            "metadata": {
                "title": blog_data['topic'],
                "description": [blog_data['intro']]
            },
            "features": [
                {
                    "title": section["title"],
                    "image_url": section.get("image_url", ""),
                    "button_text": "Try Now",
                    "button_url": "https://www.dressr.ai/clothes-swap",
                    "description": [section["description"]]
                }
                for section in blog_data['sections']
            ],
            "conclusion": [blog_data.get('conclusion', '')],
            "publish_button_text": "Publish"
        }
